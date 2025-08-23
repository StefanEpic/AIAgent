import datetime
from typing import List

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docxtpl import DocxTemplate
from langchain_gigachat.tools.giga_tool import giga_tool
from pydantic import Field

from tools.schemas import Employee


def get_current_date() -> str:
    month_ru = {
        1: "января",
        2: "февраля",
        3: "марта",
        4: "апреля",
        5: "мая",
        6: "июня",
        7: "июля",
        8: "августа",
        9: "сентября",
        10: "октября",
        11: "ноября",
        12: "декабря",
    }
    now = datetime.datetime.now()
    formatted_date = f"{now.day} {month_ru[now.month]} {now.year} г."
    return formatted_date


def set_paragraph_align(cell):
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(14)


def set_paragraph_left(cell):
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in paragraph.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(14)


create_work_permit_examples = [
    {
        "request": "создай допуск к работе для Иванова Ивана Иваныча, стажировался с 01.02.25 по 01.03.25, выходит 01.04.25",
        "params": {
            "internship_start_date": "1 февраля 2025 г.",
            "internship_end_date": "1 марта 2025 г.",
            "employees_name": "Иванов И.И.",
            "employees_return_date": "1 апреля 2025 г.",
        }
    },
    {
        "request": "надо подготовить допуск к работе для Андреевой Анны Павловны, которая проходила стажировку в период с 13 марта 2025 года по 17 апреля 2025 года. Планируемая дата выхода 23 июня 2025",
        "params": {
            "internship_start_date": "13 марта 2025 г.",
            "internship_end_date": "17 апреля 2025 г.",
            "employees_name": "Андреева А.П.",
            "employees_return_date": "23 июня 2025 г.",
        }
    },
    {
        "request": "допуск к работе борисов кирилл владимирович стажировка 05.06.25 - 14.07.25 выход на работу 30 сентября 2025",
        "params": {
            "internship_start_date": "5 июня 2025 г.",
            "internship_end_date": "14 июля 2025 г.",
            "employees_name": "Борисов К.В.",
            "employees_return_date": "30 сентября 2025 г.",
        }
    },
]


@giga_tool(few_shot_examples=create_work_permit_examples)
def create_work_permit(
        internship_start_date: str = Field(
            description="Дата начала стажировки сотрудника в формате число месяц год в формате '1 месяц 2025 г.'"),
        internship_end_date: str = Field(
            description="Дата окончания стажировки сотрудника в формате число месяц год в формате '1 месяц 2025 г.'"),
        employees_name: str = Field(description="Фамилия Имя Отчество сотрудника в формате 'Фамилия И.О.'"),
        employees_return_date: str = Field(
            description="Дата выхода на работу сотрудника в формате число месяц год в формате '1 месяц 2025 г.'")
) -> str:
    """
    Создает документ "Допуск к работе" в для определенного сотрудника.
    :param internship_start_date: Дата начала стажировки сотрудника в формате число месяц год, к примеру 1 сентября 2025 г.
    :param internship_end_date: Дата окончания стажировки сотрудника в формате число месяц год, к примеру 1 сентября 2025 г.
    :param employees_name: Фамилия Имя Отчество сотрудника в формате Фамилия И.О., к примеру Иванов И.И.
    :param employees_return_date: Дата выхода на работу сотрудника в формате число месяц год, к примеру 1 сентября 2025 г.
    :return: Ответ в виде строки
    """
    # Загрузка шаблона
    doc = DocxTemplate("./templates/start_work.docx")
    # Контекст для замены в шаблоне
    context = {
        'current_date': get_current_date(),
        'start_intern': internship_start_date,
        'finish_intern': internship_end_date,
        'person': employees_name,
        'work_start': employees_return_date,
    }
    # Рендеринг шаблона с данными
    doc.render(context)
    # Сохранение нового документа
    output_path = f"./results/Допуск к работе {employees_name}.docx"
    doc.save(output_path)
    return f"Допуск к работе для {employees_name} успешно создан!"


create_info_list_examples = [
    {
        "request": "создай лист ознакомления для Иванова Ивана Иваныча, механника 2 категории, база топлива - Центральная",
        "params": {
            "list_of_employees": [Employee(name="Иванов Иван Иванович", job_title="механник 2 категории", baza="Центральная")]
        }
    },
    {
        "request": "создай лист ознакомления для Иванова Ивана Иваныча и Петрова Константина Семеновича, сливщиков-разливщиков, база топлива - Северная",
        "params": {
            "list_of_employees": [Employee(name="Иванов Иван Иванович", job_title="сливщик-разливщик", baza="Северная"),
                                  Employee(name="Петров Константин Семенович", job_title="сливщик-разливщик", baza="Северная")]
        }
    },
    {
        "request": "создай лист ознакомления для следующих работников Иванова Ивана Иваныча главного инженера ОДКБ-5, Петрова Константина Семеновича водителя Главная",
        "params": {
            "list_of_employees": [Employee(name="Иванов Иван Иванович", job_title="главный инженер", baza="ОДКБ-5"),
                                  Employee(name="Петров Константин Семенович", job_title="водитель", baza="Главная")]
        }
    },
    {
        "request": "создай лист ознакомления для следующих бухгалтеров Иванова Ивана Иваныча, Кривоножко Светланы Геннадьевны, Смирнова Петра Степановича с Южной базы",
        "params": {
            "list_of_employees": [Employee(name="Иванов Иван Иванович", job_title="бухгалтер", baza="Южная"),
                                  Employee(name="Кривоножко Светлана Геннадьевна", job_title="бухгалтер", baza="Южная"),
                                  Employee(name="Смирнов Петр Степанович", job_title="бухгалтер", baza="Южная")]
        }
    },
]


@giga_tool(few_shot_examples=create_work_permit_examples)
def create_info_list(list_of_employees: List[Employee] = Field(description="Список сотрудников")) -> str:
    """
    Создание документа "Лист ознакомления" для нескольких сотрудников.
    :param list_of_employees: Список сотрудников, в которых содержится ФИО сотрудника, должность и его топливная база
    :return: Ответ в виде строки
    """
    doc = Document("./templates/info_list.docx")
    table = doc.tables[0]

    for row_num, row_data in enumerate(list_of_employees, 1):
        row = table.add_row().cells
        # Вставка данных
        row[0].text = str(row_num)
        row[1].text = row_data.name
        row[2].text = row_data.job_title
        row[3].text = row_data.baza
        # Оформление
        row[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        row[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        row[2].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        row[3].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_paragraph_align(row[0])
        set_paragraph_left(row[1])
        set_paragraph_align(row[2])
        set_paragraph_align(row[3])

    doc.save("./results/Лист ознакомления.docx")
    return "Лист ознакомления успешно создан!"
